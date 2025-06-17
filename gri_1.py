import streamlit as st
import pandas as pd
import datetime as dt
import json
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt

try:
    from docx2pdf import convert as docx2pdf_convert
except ImportError:
    docx2pdf_convert = None


# Define emission factors
INDIA_SCOPE1_EF = {
    "Diesel (litres)": 2.68,
    "Petrol (litres)": 2.31,
    "Furnace Oil (litres)": 3.10,
    "LPG (kg)": 2.95,
    "Natural Gas (scm)": 1.93,
}
INDIA_GRID_EF_T_PER_KWH = 0.00082
DEFRA_SCOPE3_EF = {
    "Purchased Goods & Services (₹ lakh)": 1.95,
    "Capital Goods (₹ lakh)": 2.40,
    "Fuel & Energy‑Related Activities (GJ)": 0.000125,
    "Upstream T&D (t‑km)": 0.000055,
    "Waste Generated (t)": 0.98,
    "Business Travel (passenger‑km)": 0.00018,
    "Employee Commuting (passenger‑km)": 0.00012,
    "End‑of‑Life Treatment (t)": 1.15,
}
MATERIAL_TOPICS = [
    "Employee Health, Safety & Wellbeing", "Climate Action", "Product Stewardship",
    "Responsible Supply Chain", "Talent Management & Training", "Human Rights & Labour Practices",
    "Circular Economy", "Environmental Protection", "Business Ethics", "Corruption",
    "Air Pollution", "Water", "Community Relations", "Data Privacy",
]

def ai_narrative(topic, info):
    return "(AI narrative disabled: OpenAI not configured)"

def build_docx(data):
    doc = Document()
    doc.add_heading("GRI ESG Report – COPMANY NAME", 0)
    doc.add_paragraph(f"Reporting period: {data['period']}")
    doc.add_paragraph(f"Frameworks: {', '.join(data['frameworks'])}")
    doc.add_heading("Material Topic Narratives", level=1)
    for topic, txt in data.get("narratives", {}).items():
        doc.add_heading(topic, level=2)
        doc.add_paragraph(txt)
    doc.add_heading("Environmental KPIs", level=1)
    for k, v in data["environmental"].items():
        doc.add_paragraph(f"{k}: {v}")
    doc.add_heading("Scope 3 Breakdown", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Category", "Qty", "EF", "Emissions"
    for row in data["scope3_details"]:
        r = table.add_row().cells
        r[0].text = row["Category"]
        r[1].text = str(row["Qty"])
        r[2].text = str(row["EF"])
        r[3].text = f"{row['Emissions']:.2f}"
    buf = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(buf.name)
    with open(buf.name, "rb") as f:
        return f.read()

st.set_page_config(page_title="GRI ESG Reporter", layout="wide")
st.title("🌱 COMAPANY GRI Report")

company = st.text_input("Company Name", "COMPANY")
report_period = st.text_input("Reporting Period", "01 April 2025 – 31 March 2026")
frameworks = st.multiselect("Frameworks Used", ["GRI Standards 2021", "BRSR", "SDGs"], default=["GRI Standards 2021", "SDGs"])

st.subheader("Material Topics")
selected_topics = st.multiselect("Choose topics", MATERIAL_TOPICS, default=MATERIAL_TOPICS)
topic_info, narratives = {}, {}
for topic in selected_topics:
    with st.expander(topic):
        stk = st.text_input("Stakeholders", key=topic+"stk")
        rsk = st.text_area("Risks", key=topic+"rsk")
        opp = st.text_area("Opportunities", key=topic+"opp")
        kpis = [x.strip() for x in st.text_area("KPIs (comma-separated)", key=topic+"kpi").split(",") if x.strip()]
        topic_info[topic] = {"stakeholders": stk, "risks": rsk, "opportunities": opp, "kpis": kpis}

if st.checkbox("Generate AI narratives (optional)", value=False):
    for t, info in topic_info.items():
        narratives[t] = ai_narrative(t, info)

st.subheader("Scope 1 Emissions (India Fuel-based)")
scope1_total = 0.0
for fuel, ef in INDIA_SCOPE1_EF.items():
    qty = st.number_input(f"{fuel} used", 0.0, key=fuel)
    scope1_total += qty * ef
st.success(f"Total Scope 1: {scope1_total:.2f} tCO₂e")

st.subheader("Scope 2 Emissions (CEA Grid)")
kwh = st.number_input("Electricity purchased (kWh)", 0.0)
scope2_total = kwh * INDIA_GRID_EF_T_PER_KWH
st.success(f"Total Scope 2: {scope2_total:.2f} tCO₂e")

st.subheader("Scope 3 (DEFRA)")
scope3_rows, total_scope3 = [], 0.0
for cat, ef in DEFRA_SCOPE3_EF.items():
    with st.expander(cat):
        qty = st.number_input("Quantity", 0.0, key=cat)
        ef_input = st.number_input("Emission Factor", ef, key=cat+"_ef")
        em = qty * ef_input
        total_scope3 += em
        scope3_rows.append({"Category": cat, "Qty": qty, "EF": ef_input, "Emissions": em})
st.success(f"Total Scope 3: {total_scope3:.2f} tCO₂e")

st.subheader("Environmental Data")
energy_gj = st.number_input("Energy (GJ)", 0.0)
renew_gj = st.number_input("Renewable Energy (GJ)", 0.0)
water_m3 = st.number_input("Water Withdrawal (m³)", 0.0)
waste_t = st.number_input("Waste Generated (t)", 0.0)

st.subheader("Social KPIs")
emp = st.number_input("Total Employees", 0)
ltifr = st.number_input("LTIFR", 0.0)
train_hr = st.number_input("Training Hours", 0.0)
female_pct = st.number_input("% Female Employees", 0.0)

st.subheader("Governance KPIs")
ethics_pct = st.number_input("% Trained in Ethics", 0.0)
board_size = st.number_input("Board Size", 0)
indep_dir = st.number_input("Independent Directors", 0)
data_breach = st.number_input("Data Breaches", 0)

report = {
    "company": company,
    "period": report_period,
    "frameworks": frameworks,
    "topics": topic_info,
    "narratives": narratives,
    "environmental": {
        "Scope 1": round(scope1_total, 2),
        "Scope 2": round(scope2_total, 2),
        "Scope 3": round(total_scope3, 2),
        "Energy": energy_gj,
        "Renewable": renew_gj,
        "Water": water_m3,
        "Waste": waste_t,
    },
    "scope3_details": scope3_rows,
    "social": {
        "Employees": emp,
        "LTIFR": ltifr,
        "Training Hours": train_hr,
        "% Female": female_pct,
    },
    "governance": {
        "% Ethics Trained": ethics_pct,
        "Data Breaches": data_breach,
        "Board Size": board_size,
        "Independent Directors": indep_dir,
    },
    "created": str(dt.datetime.now()),
}

st.subheader("Export Report")

# JSON Export
st.download_button("⬇ Download JSON", json.dumps(report, indent=2),
                   file_name="gri_tmseating_report.json", mime="application/json")

# Word Export (RECOMMENDED)
docx_bytes = build_docx(report)
st.download_button("⬇ Download Word (.docx)", docx_bytes,
                   file_name="gri_tmseating_report.docx",
                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.caption("💡 To export as PDF, open the Word file and choose **File → Save As → PDF** in Microsoft Word.")

st.markdown("---")
st.caption("Scientific fact: LPG (liquefied petroleum gas) has a lower CO₂ emission factor per unit of energy than coal, making it a transitional fuel in many emission reduction policies.")

